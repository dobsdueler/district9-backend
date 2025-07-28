from flask import Flask, request, send_file, jsonify, render_template
from flask_cors import CORS
from docx import Document
import os
import io
import smtplib
from email.message import EmailMessage
from dotenv import load_dotenv
import datetime

load_dotenv()

app = Flask(__name__)
CORS(app)

# === Function to generate Word document ===
def generate_word_document(data):
    doc = Document()
    doc.add_heading("Classroom Observation Report", level=1)

    doc.add_paragraph(f"Teacher: {data.get('teacher_name')}")
    doc.add_paragraph(f"School: {data.get('school')}")
    doc.add_paragraph(f"Observed By: {data.get('observed_by')}")
    doc.add_paragraph(f"Indicator: {data.get('indicator')}")
    doc.add_paragraph(f"Rating: {data.get('rating')}")
    doc.add_paragraph(f"Notes: {data.get('notes')}")
    doc.add_paragraph(f"Date: {data.get('date')}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# === Function to send email ===
def send_email_report(to_email, subject, body, doc_buffer, filename="report.docx"):
    email_address = os.environ.get("dobsdueler@gmail.com")
    email_password = os.environ.get("yniu lxrv pawt wkix")

    if not email_address or not email_password:
        return {"status": "error", "message": "Email credentials not set."}

    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = email_address
    msg["To"] = to_email
    msg.set_content(body)

    doc_buffer.seek(0)
    msg.add_attachment(doc_buffer.read(), maintype="application",
                       subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
                       filename=filename)

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
        smtp.login(email_address, email_password)
        smtp.send_message(msg)

# === ROUTES ===

@app.route("/")
def index():
    return render_template('index.html')  # Now renders the index.html form

@app.route("/submit-report", methods=["POST"])
def submit_report():
    school = request.form['school']
    summary = request.form['summary']
    print(f"Received report from {school}: {summary}")
    return f"Report from {school} received!"

@app.route("/observe-class", methods=["POST"])
def observe_class():
    data = request.form
    doc_buffer = generate_word_document(data)

    return send_file(
        doc_buffer,
        as_attachment=True,
        download_name="Observation_Report.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.route("/api/send-report", methods=["POST"])
def send_report():
    data = request.json
    recipient_email = data.get("email")
    subject = data.get("subject", "District 9 Observation Report")
    content = data.get("content", "Attached is the report for your reference.")
    observations = data.get("observations", [])

    if not recipient_email or not observations:
        return jsonify({"error": "Missing email or observations"}), 400

    doc = Document()
    doc.add_heading('District 9 Observation Report', 0)

    for obs in observations:
        doc.add_paragraph(f"Teacher: {obs['teacher_name']}")
        doc.add_paragraph(f"School: {obs['school']}")
        doc.add_paragraph(f"Observed by: {obs['observed_by']}")
        doc.add_paragraph(f"Indicator: {obs['indicator']}")
        doc.add_paragraph(f"Rating: {obs['rating']}")
        doc.add_paragraph(f"Notes: {obs['notes']}")
        doc.add_paragraph(f"Date: {obs['date']}")
        doc.add_paragraph("\n---\n")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    try:
        email_user = os.getenv("EMAIL_USER")
        email_pass = os.getenv("EMAIL_PASS")

        msg = EmailMessage()
        msg['Subject'] = subject
        msg['From'] = email_user
        msg['To'] = recipient_email
        msg.set_content(content)

        filename = f"District9_Report_{datetime.datetime.now().strftime('%Y%m%d')}.docx"
        msg.add_attachment(buffer.read(), maintype='application',
                           subtype='vnd.openxmlformats-officedocument.wordprocessingml.document',
                           filename=filename)

        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
            smtp.login(email_user, email_pass)
            smtp.send_message(msg)

        return jsonify({"message": "Email sent successfully"}), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5050)
